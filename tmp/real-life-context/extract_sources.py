from __future__ import annotations

import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader


ROOT = Path(
    "/Users/juliocoraspe/Library/CloudStorage/GoogleDrive-"
    "juliocoraspe@gmail.com/My Drive/BeReal Externship"
)
OUT = Path("/Users/juliocoraspe/Desktop/UX UI Portfolio/tmp/real-life-context/extracted")


DOCX_FILES = [
    ROOT / "project 2/BeReal_Retention_Analysis.docx",
    ROOT / "project 2/Feature_Consolidation.docx",
    ROOT / "project 3/BeReal_Feature_Ideas_Table.docx",
    ROOT / "project 3/Challenges_Brainstorm.docx",
    ROOT / "project 3/PRD_Lite.docx",
    ROOT / "project 3/RLChallenges_PRD_Lite.docx",
    ROOT / "project 3/User_Stories_Acceptance_Criteria.docx",
    ROOT / "project 4/focus-group-plan.docx",
    ROOT / "project 4/discussion-guide.docx",
    ROOT / "project 4/test_summary_report_bereal.docx",
]

PDF_FILES = [
    ROOT / "project 1/BeReal — Research Analysis.pdf",
    ROOT / "project 1/BeReal — PM Problem Statement & Opportunity Note.pdf",
    ROOT / "project 3/BeReal_RICE_Analysis_Real_Life_Challenges.pdf",
    ROOT / "project 3/Real-Life Challenges — BeReal Feature Proposal.pdf",
]

TEXT_FILES = [
    ROOT / "project 4/transcript_focus_group_real_life_challenges.txt",
    ROOT / "PROJECT_4_SYNTHESIS_FocusGroup.md",
    ROOT / "CAUSAL_THREAD_Problem-to-Solution.md",
    ROOT / "NARRATIVE_ARC_presentation.md",
    ROOT / "PRESENTATION_STRUCTURE_Real-Life-Challenges.md",
    ROOT / "PITCH_GAP_ANALYSIS_and_7-Slide-Mapping.md",
    ROOT / "ANTICIPATED_QUESTIONS_NextSteps_Limitations.md",
]


def safe_name(path: Path) -> str:
    value = path.relative_to(ROOT).as_posix()
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value)
    return value.strip("_")


def extract_docx(path: Path) -> str:
    doc = Document(path)
    lines: list[str] = [f"SOURCE: {path}", ""]
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            lines.append(f"[{style}] {text}" if style else text)
    for index, table in enumerate(doc.tables, start=1):
        lines.extend(["", f"TABLE {index}"])
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", cell.text.strip()) for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf(path: Path) -> str:
    reader = PdfReader(path)
    lines = [f"SOURCE: {path}", f"PAGES: {len(reader.pages)}", ""]
    for page_number, page in enumerate(reader.pages, start=1):
        lines.extend([f"--- PAGE {page_number} ---", page.extract_text() or ""])
    return "\n".join(lines)


def extract_workbook(path: Path) -> dict[str, list[dict[str, object]]]:
    workbook = pd.ExcelFile(path)
    sheets: dict[str, list[dict[str, object]]] = {}
    for sheet_name in workbook.sheet_names:
        frame = pd.read_excel(path, sheet_name=sheet_name, dtype=object)
        frame = frame.where(pd.notna(frame), None)
        sheets[sheet_name] = frame.to_dict(orient="records")
    return sheets


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    manifest: list[dict[str, object]] = []

    for path in DOCX_FILES:
        if not path.exists():
            manifest.append({"path": str(path), "status": "missing"})
            continue
        output = OUT / f"{safe_name(path)}.txt"
        output.write_text(extract_docx(path), encoding="utf-8")
        manifest.append({"path": str(path), "output": str(output), "status": "ok"})

    for path in PDF_FILES:
        if not path.exists():
            manifest.append({"path": str(path), "status": "missing"})
            continue
        output = OUT / f"{safe_name(path)}.txt"
        output.write_text(extract_pdf(path), encoding="utf-8")
        manifest.append({"path": str(path), "output": str(output), "status": "ok"})

    for path in TEXT_FILES:
        if not path.exists():
            manifest.append({"path": str(path), "status": "missing"})
            continue
        output = OUT / f"{safe_name(path)}.txt"
        output.write_text(path.read_text(encoding="utf-8", errors="replace"), encoding="utf-8")
        manifest.append({"path": str(path), "output": str(output), "status": "ok"})

    workbook = ROOT / "project 4/Real Life Challenge (Responses).xlsx"
    if workbook.exists():
        output = OUT / "project_4_Real_Life_Challenge_Responses.json"
        output.write_text(
            json.dumps(extract_workbook(workbook), ensure_ascii=False, indent=2, default=str),
            encoding="utf-8",
        )
        manifest.append({"path": str(workbook), "output": str(output), "status": "ok"})
    else:
        manifest.append({"path": str(workbook), "status": "missing"})

    (OUT / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()
